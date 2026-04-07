"""
角色扮演 API 路由
"""

import os
import re
import asyncio
import tempfile
from pathlib import Path
from typing import List, Optional
from datetime import datetime

from fastapi import APIRouter, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse
from pydantic import BaseModel

from character_analyzer import CharacterAnalyzer, CharacterProfile
from character_db import db


router = APIRouter(prefix="/api/roleplay", tags=["角色扮演"])

# 全局分析器实例
_analyzer = None


# ==================== 文件解析工具 ====================

def extract_text_from_file(content: bytes, file_ext: str) -> str:
    """
    从各种格式的文件中提取文本
    支持: txt, pdf, docx, epub, md, html
    """
    text = ""
    
    try:
        if file_ext == '.txt':
            # 尝试多种编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    text = content.decode(encoding)
                    break
                except:
                    continue
            if not text:
                text = content.decode('utf-8', errors='ignore')
                
        elif file_ext == '.pdf':
            text = extract_text_from_pdf(content)
            
        elif file_ext == '.docx':
            text = extract_text_from_docx(content)
            
        elif file_ext == '.doc':
            # 尝试用 pandoc 或者直接读取
            text = extract_text_from_doc(content)
            
        elif file_ext == '.epub':
            text = extract_text_from_epub(content)
            
        elif file_ext in ['.md', '.markdown']:
            text = content.decode('utf-8', errors='ignore')
            
        elif file_ext == '.html':
            text = extract_text_from_html(content)
            
    except Exception as e:
        print(f"文件解析错误: {e}")
        text = ""
    
    return text.strip()


def extract_text_from_pdf(content: bytes) -> str:
    """从 PDF 提取文本，优先使用文本层，降级使用 OCR"""
    import pdfplumber
    
    text = ""
    has_text = False
    
    # 方法1：尝试直接提取文本层
    try:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        with pdfplumber.open(tmp_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        os.unlink(tmp_path)
        
        # 检查是否提取到足够文本（文本层可能有内容）
        if len(text.strip()) > 100:
            has_text = True
            
    except Exception as e:
        print(f"PDF 文本提取失败: {e}")
    
    # 方法2：如果文本层内容太少，尝试 OCR
    if not has_text or len(text.strip()) < 50:
        print("PDF 文本层内容过少，尝试 OCR...")
        ocr_text = extract_text_from_pdf_ocr(content)
        if ocr_text and len(ocr_text.strip()) > len(text.strip()):
            text = ocr_text
    
    return text


def extract_text_from_pdf_ocr(content: bytes) -> str:
    """使用 OCR 从 PDF 提取文本（影印版 PDF）"""
    try:
        import pytesseract
        from PIL import Image
        from pdf2image import convert_from_bytes
        
        # 将 PDF 转换为图片
        images = convert_from_bytes(content, dpi=200)
        
        text = ""
        for i, img in enumerate(images):
            print(f"OCR 处理第 {i+1}/{len(images)} 页...")
            page_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            text += f"[第{i+1}页]\n{page_text}\n"
        
        return text
        
    except ImportError as e:
        print(f"OCR 依赖未安装: {e}")
        print("提示: 运行 'pip install pytesseract pillow pdf2image' 并安装 Tesseract OCR")
        return ""
    except Exception as e:
        print(f"OCR 处理失败: {e}")
        return ""


def extract_text_from_docx(content: bytes) -> str:
    """从 Word DOCX 提取文本"""
    try:
        from docx import Document
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        doc = Document(tmp_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text += "\n" + row_text
        
        os.unlink(tmp_path)
        return text
        
    except Exception as e:
        print(f"DOCX 解析失败: {e}")
        return ""


def extract_text_from_doc(content: bytes) -> str:
    """从旧版 Word DOC 提取文本（需要安装 antiword 或 LibreOffice）"""
    # 尝试用 python-docx2txt 或其他方式
    # 如果安装了 LibreOffice 可以转换
    try:
        import subprocess
        # 尝试调用 LibreOffice
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'txt', '--outdir', tempfile.gettempdir()],
            input=content,
            capture_output=True,
            timeout=30
        )
        return ""
    except:
        return ""


def extract_text_from_epub(content: bytes) -> str:
    """从 EPUB 提取文本"""
    try:
        import zipfile
        from xml.etree import ElementTree
        
        with tempfile.NamedTemporaryFile(suffix='.epub', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        text = ""
        with zipfile.ZipFile(tmp_path, 'r') as epub:
            # 找到所有 HTML/XHTML 文件
            for name in epub.namelist():
                if name.endswith(('.html', '.xhtml', '.htm')):
                    try:
                        html_content = epub.read(name).decode('utf-8', errors='ignore')
                        # 简单提取文本
                        text += extract_text_from_html(html_content.encode()) + "\n"
                    except:
                        continue
        
        os.unlink(tmp_path)
        return text
        
    except Exception as e:
        print(f"EPUB 解析失败: {e}")
        return ""


def extract_text_from_html(content: bytes) -> str:
    """从 HTML 提取纯文本"""
    try:
        from bs4 import BeautifulSoup
        
        soup = BeautifulSoup(content, 'html.parser')
        
        # 移除脚本和样式
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text()
        
        # 清理空白
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return text
        
    except Exception as e:
        # 降级：简单正则提取
        html = content.decode('utf-8', errors='ignore')
        text = re.sub(r'<[^>]+>', ' ', html)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()


def get_analyzer() -> CharacterAnalyzer:
    """获取或创建分析器实例"""
    global _analyzer
    if _analyzer is None:
        api_key = os.getenv("DEEPSEEK_API_KEY", "")
        if not api_key:
            raise HTTPException(status_code=500, detail="DeepSeek API Key 未配置")
        _analyzer = CharacterAnalyzer(api_key)
    return _analyzer


# ==================== 请求/响应模型 ====================

class QuestionRequest(BaseModel):
    character_ids: List[str]
    question: str
    mode: str = "parallel"  # parallel 或 debate


class DebateRequest(BaseModel):
    character_id_1: str
    character_id_2: str
    question: str


class CreateGroupRequest(BaseModel):
    name: str
    description: str = ""
    character_ids: List[str] = []


class UpdateCharacterRequest(BaseModel):
    name: Optional[str] = None
    tags: Optional[List[str]] = None
    personality_summary: Optional[str] = None


# ==================== 角色库 API ====================

@router.get("/characters")
async def list_characters():
    """获取所有角色"""
    characters = db.get_all_characters()
    return {
        "code": 0,
        "data": {
            "characters": [c.to_dict() for c in characters],
            "total": len(characters)
        }
    }


@router.get("/characters/{character_id}")
async def get_character(character_id: str):
    """获取单个角色详情"""
    character = db.get_character(character_id)
    if not character:
        raise HTTPException(status_code=404, detail="角色不存在")
    
    return {
        "code": 0,
        "data": character.to_dict()
    }


@router.delete("/characters/{character_id}")
async def delete_character(character_id: str):
    """删除角色"""
    if db.delete_character(character_id):
        return {"code": 0, "message": "删除成功"}
    raise HTTPException(status_code=404, detail="角色不存在")


@router.put("/characters/{character_id}")
async def update_character(character_id: str, request: UpdateCharacterRequest):
    """更新角色信息"""
    updates = {k: v for k, v in request.model_dump().items() if v is not None}
    character = db.update_character(character_id, updates)
    if character:
        return {"code": 0, "data": character.to_dict()}
    raise HTTPException(status_code=404, detail="角色不存在")


@router.get("/characters/search")
async def search_characters(q: str = ""):
    """搜索角色"""
    if not q:
        return {"code": 0, "data": {"characters": [], "total": 0}}
    
    results = db.search_characters(q)
    return {
        "code": 0,
        "data": {
            "characters": [c.to_dict() for c in results],
            "total": len(results)
        }
    }


# ==================== 书籍上传与角色提取 ====================

@router.post("/upload-book")
async def upload_book(
    file: UploadFile = File(...),
    title: str = Form(""),
    author: str = Form("")
):
    """上传书籍并提取角色"""
    # 支持的文件格式
    allowed_types = ['.txt', '.pdf', '.docx', '.doc', '.epub', '.md', '.markdown', '.html']
    file_ext = Path(file.filename).suffix.lower()
    
    if file_ext not in allowed_types:
        raise HTTPException(
            status_code=400, 
            detail=f"不支持的文件类型。支持的格式: {', '.join(allowed_types)}"
        )
    
    # 读取文件内容
    content = await file.read()
    
    print(f"📖 正在解析文件: {file.filename} ({len(content)} bytes)")
    
    # 提取文本
    text = extract_text_from_file(content, file_ext)
    
    if not text or len(text.strip()) < 100:
        raise HTTPException(
            status_code=400, 
            detail="文本内容过短或无法解析，无法提取角色。\n提示：如果是影印版 PDF，系统将尝试 OCR 识别，请确保 Tesseract OCR 已安装。"
        )
    
    print(f"✅ 提取到 {len(text)} 字符")
    
    # 提取书名
    if not title and file.filename:
        title = Path(file.filename).stem
    
    # 使用 AI 分析角色
    analyzer = get_analyzer()
    
    try:
        characters = await analyzer.extract_characters(
            text=text,
            book_title=title,
            book_author=author
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"角色分析失败: {str(e)}")
    
    # 保存书籍信息
    book = db.add_book(title=title, author=author, file_path=file.filename)
    
    # 保存角色到数据库
    saved_characters = []
    for char in characters:
        char.book_title = title  # 确保关联正确的书籍
        saved = db.add_character(char)
        saved_characters.append(saved)
    
    # 更新书籍的角色数量
    db.update_book_character_count(book["id"])
    
    return {
        "code": 0,
        "data": {
            "book": book,
            "characters": [c.to_dict() for c in saved_characters],
            "total": len(saved_characters),
            "total_characters": len(saved_characters)
        }
    }


@router.post("/extract-characters")
async def extract_from_text(
    text: str = Form(...),
    book_title: str = Form(""),
    book_author: str = Form("")
):
    """从文本直接提取角色"""
    if len(text) < 100:
        raise HTTPException(status_code=400, detail="文本内容过短")
    
    analyzer = get_analyzer()
    
    try:
        characters = await analyzer.extract_characters(
            text=text,
            book_title=book_title or "自定义文本",
            book_author=book_author
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"角色分析失败: {str(e)}")
    
    # 保存角色
    saved_characters = []
    for char in characters:
        saved = db.add_character(char)
        saved_characters.append(saved)
    
    return {
        "code": 0,
        "data": {
            "characters": [c.to_dict() for c in saved_characters],
            "total": len(saved_characters)
        }
    }


# ==================== 角色组管理 ====================

@router.get("/groups")
async def list_groups():
    """获取所有角色组"""
    groups = db.get_all_groups()
    
    # 补充每个组的角色信息
    for group in groups:
        chars = db.get_group_characters(group["id"])
        group["characters"] = [c.to_dict() for c in chars]
        group["character_count"] = len(chars)
    
    return {
        "code": 0,
        "data": groups
    }


@router.post("/groups")
async def create_group(request: CreateGroupRequest):
    """创建角色组"""
    group = db.create_group(
        name=request.name,
        description=request.description,
        character_ids=request.character_ids
    )
    
    chars = db.get_group_characters(group["id"])
    group["characters"] = [c.to_dict() for c in chars]
    
    return {"code": 0, "data": group}


@router.delete("/groups/{group_id}")
async def delete_group(group_id: str):
    """删除角色组"""
    if db.delete_group(group_id):
        return {"code": 0, "message": "删除成功"}
    raise HTTPException(status_code=404, detail="角色组不存在")


@router.post("/groups/{group_id}/characters/{character_id}")
async def add_to_group(group_id: str, character_id: str):
    """添加角色到组"""
    if db.add_to_group(group_id, character_id):
        return {"code": 0, "message": "添加成功"}
    raise HTTPException(status_code=404, detail="角色组不存在")


@router.delete("/groups/{group_id}/characters/{character_id}")
async def remove_from_group(group_id: str, character_id: str):
    """从组中移除角色"""
    if db.remove_from_group(group_id, character_id):
        return {"code": 0, "message": "移除成功"}
    raise HTTPException(status_code=404, detail="角色组不存在")


# ==================== 难题解答 ====================

@router.post("/ask")
async def ask_character(request: QuestionRequest):
    """向选中的角色提问"""
    if not request.character_ids:
        raise HTTPException(status_code=400, detail="请至少选择一个角色")
    
    if not request.question:
        raise HTTPException(status_code=400, detail="请输入问题")
    
    # 获取角色
    characters = []
    for char_id in request.character_ids:
        char = db.get_character(char_id)
        if char:
            characters.append(char)
    
    if not characters:
        raise HTTPException(status_code=404, detail="未找到选中的角色")
    
    analyzer = get_analyzer()
    
    if request.mode == "debate" and len(characters) >= 2:
        # 辩论模式
        debate = await analyzer.generate_debate(
            characters[0], 
            characters[1], 
            request.question
        )
        return {
            "code": 0,
            "data": {
                "mode": "debate",
                "question": request.question,
                "debate": debate
            }
        }
    else:
        # 并行回答模式
        results = await analyzer.generate_group_response(characters, request.question)
        
        return {
            "code": 0,
            "data": {
                "mode": "parallel",
                "question": request.question,
                "responses": results
            }
        }


@router.post("/debate")
async def create_debate(request: DebateRequest):
    """创建两个角色的辩论"""
    char1 = db.get_character(request.character_id_1)
    char2 = db.get_character(request.character_id_2)
    
    if not char1 or not char2:
        raise HTTPException(status_code=404, detail="角色不存在")
    
    analyzer = get_analyzer()
    
    debate = await analyzer.generate_debate(char1, char2, request.question)
    
    return {
        "code": 0,
        "data": debate
    }


# ==================== 书籍管理 ====================

@router.get("/books")
async def list_books():
    """获取所有书籍"""
    books = db.get_all_books()
    
    # 为每本书补充角色信息
    for book in books:
        chars = db.get_characters_by_book(book["title"])
        book["characters"] = [c.to_dict() for c in chars]
        book["character_count"] = len(chars)
    
    return {
        "code": 0,
        "data": books
    }


@router.delete("/books/{book_id}")
async def delete_book(book_id: str, delete_characters: bool = False):
    """删除书籍"""
    if db.delete_book(book_id, delete_characters):
        return {"code": 0, "message": "删除成功"}
    raise HTTPException(status_code=404, detail="书籍不存在")


# ==================== 预设角色 ====================

@router.get("/presets")
async def list_presets():
    """获取预设角色库"""
    # 这里可以添加一些经典角色作为预设
    presets = [
        {
            "id": "preset_sherlock",
            "name": "夏洛克·福尔摩斯",
            "source": "福尔摩斯探案集",
            "personality": "理性、逻辑、冷静、观察力强、略带傲慢",
            "tags": ["侦探", "推理", "经典"]
        },
        {
            "id": "preset_watson",
            "name": "约翰·华生",
            "source": "福尔摩斯探案集",
            "personality": "忠诚、务实、有人情味、略显保守",
            "tags": ["医生", "侦探", "经典"]
        },
        {
            "id": "preset_poirot",
            "name": "赫尔克里·波洛",
            "source": "阿加莎·克里斯蒂",
            "personality": "挑剔、精确、道德感强、喜欢秩序",
            "tags": ["侦探", "比利时", "经典"]
        },
        {
            "id": "preset_marple",
            "name": "简·马普尔",
            "source": "阿加莎·克里斯蒂",
            "personality": "慈祥、敏锐、洞察人性、了解人性弱点",
            "tags": ["侦探", "英国乡村", "经典"]
        }
    ]
    
    return {
        "code": 0,
        "data": presets
    }
